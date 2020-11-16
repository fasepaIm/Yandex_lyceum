from docx import Document
from docx.shared import Inches
import sys

data = sys.stdin.readlines()

da = data[:2]
data = data[2:]

document = Document()

for i in data:
    if data.index(i) != 0:
        document.add_page_break()
    document.add_heading('Приглашение', 0)
    document.add_paragraph(f'{i[:-2]},')

    p = document.add_paragraph(f'{da[0][:-2].lower().capitalize()}\
{da[-1][:-2].lower()} состоится мероприятие. ')
    p.add_run('Приглашаем вас посетить его.')

document.save('test.docx')